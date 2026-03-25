#!/usr/bin/env python3
"""
Dataset Conversion Tool for AnonLFI Benchmark

Converts vulnnet_scans_openvas dataset files to multiple formats:
- CSV → XLSX (Excel)
- CSV → DOCX (Word table)
- CSV/XML → JSON (structured data)
- PDF → Images (one per page, for OCR testing)

Features:
- Modular, extensible converter architecture
- Robust error handling with detailed logging
- Parallel processing support
- Progress tracking
- Preserves original data structure
- SOLID principles and type safety

Author: AnonShield Team
Date: February 2026
"""

import argparse
import csv
import json
import logging
import shutil
import sys
import xml.etree.ElementTree as ET
from abc import ABC, abstractmethod
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Any
from datetime import datetime

# Third-party imports (will be checked at runtime)
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False


# =============================================================================
# CONFIGURATION & DATA CLASSES
# =============================================================================

@dataclass
class ConversionConfig:
    """Configuration for dataset conversion."""
    source_dir: Path
    output_dir: Path
    formats: Set[str]
    max_workers: int = 4
    skip_existing: bool = True
    image_dpi: int = 150
    image_format: str = "png"
    verbose: bool = False
    
    def __post_init__(self):
        """Validate configuration."""
        if not self.source_dir.exists():
            raise ValueError(f"Source directory does not exist: {self.source_dir}")
        
        valid_formats = {"xlsx", "docx", "json", "pdf_images"}
        invalid = self.formats - valid_formats
        if invalid:
            raise ValueError(f"Invalid formats: {invalid}. Valid: {valid_formats}")


@dataclass
class ConversionResult:
    """Result of a single file conversion."""
    source_file: Path
    target_format: str
    output_file: Optional[Path] = None
    success: bool = False
    error: Optional[str] = None
    duration_sec: float = 0.0
    output_size_bytes: int = 0


@dataclass
class ConversionSummary:
    """Summary of all conversions."""
    total_files: int = 0
    successful: int = 0
    failed: int = 0
    skipped: int = 0
    results: List[ConversionResult] = field(default_factory=list)
    start_time: datetime = field(default_factory=datetime.now)
    end_time: Optional[datetime] = None
    
    def add_result(self, result: ConversionResult):
        """Add a conversion result and update counters."""
        self.results.append(result)
        if result.success:
            self.successful += 1
        elif result.error and "skipped" in result.error.lower():
            self.skipped += 1
        else:
            self.failed += 1
    
    def finalize(self):
        """Mark conversion as complete."""
        self.end_time = datetime.now()
    
    @property
    def duration_sec(self) -> float:
        """Total duration in seconds."""
        if self.end_time:
            return (self.end_time - self.start_time).total_seconds()
        return 0.0


# =============================================================================
# CONVERTER INTERFACE & IMPLEMENTATIONS
# =============================================================================

class FileConverter(ABC):
    """Abstract base class for file converters."""
    
    def __init__(self, config: ConversionConfig, logger: logging.Logger):
        self.config = config
        self.logger = logger
    
    @abstractmethod
    def can_convert(self, source_file: Path) -> bool:
        """Check if this converter can handle the source file."""
        pass
    
    @abstractmethod
    def get_output_path(self, source_file: Path) -> Path:
        """Generate output file path."""
        pass
    
    @abstractmethod
    def convert(self, source_file: Path) -> ConversionResult:
        """Convert the file and return result."""
        pass
    
    def _should_skip(self, source_file: Path, output_file: Path) -> bool:
        """Check if conversion should be skipped."""
        if not self.config.skip_existing:
            return False
        
        if output_file.exists():
            # Skip if output is newer and not empty
            if (output_file.stat().st_mtime >= source_file.stat().st_mtime and
                output_file.stat().st_size > 0):
                return True
        
        return False


class CSVToXLSXConverter(FileConverter):
    """Converts CSV files to XLSX format using openpyxl."""
    
    TARGET_FORMAT = "xlsx"
    
    def can_convert(self, source_file: Path) -> bool:
        return source_file.suffix.lower() == ".csv" and OPENPYXL_AVAILABLE
    
    def get_output_path(self, source_file: Path) -> Path:
        """Generate output path maintaining directory structure."""
        rel_path = source_file.relative_to(self.config.source_dir)
        output_path = self.config.output_dir / self.TARGET_FORMAT / rel_path.parent / f"{source_file.stem}.xlsx"
        return output_path
    
    def convert(self, source_file: Path) -> ConversionResult:
        """Convert CSV to XLSX with formatting."""
        start_time = datetime.now()
        result = ConversionResult(
            source_file=source_file,
            target_format=self.TARGET_FORMAT
        )
        
        try:
            output_file = self.get_output_path(source_file)
            
            # Check if should skip
            if self._should_skip(source_file, output_file):
                result.error = f"Skipped: output exists and is up-to-date"
                result.output_file = output_file
                return result
            
            # Create output directory
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Read CSV and write to XLSX
            wb = Workbook()
            ws = wb.active
            ws.title = "Scan Results"
            
            with open(source_file, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                
                # Write header with formatting
                header = next(reader)
                ws.append(header)
                
                # Format header row
                header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                header_font = Font(bold=True, color="FFFFFF")
                
                for cell in ws[1]:
                    cell.fill = header_fill
                    cell.font = header_font
                
                # Write data rows
                for row in reader:
                    ws.append(row)
            
            # Auto-adjust column widths (basic implementation)
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)  # Cap at 50
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # Save workbook
            wb.save(output_file)
            
            result.success = True
            result.output_file = output_file
            result.output_size_bytes = output_file.stat().st_size
            
            self.logger.info(f"✓ CSV → XLSX: {source_file.name} ({result.output_size_bytes / 1024:.1f} KB)")
            
        except Exception as e:
            result.error = str(e)
            self.logger.error(f"✗ CSV → XLSX failed: {source_file.name}: {e}")
        
        finally:
            result.duration_sec = (datetime.now() - start_time).total_seconds()
        
        return result


class TXTToDOCXConverter(FileConverter):
    """Converts TXT files to DOCX format."""
    
    TARGET_FORMAT = "docx"
    
    def can_convert(self, source_file: Path) -> bool:
        return source_file.suffix.lower() == ".txt" and DOCX_AVAILABLE
    
    def get_output_path(self, source_file: Path) -> Path:
        rel_path = source_file.relative_to(self.config.source_dir)
        output_path = self.config.output_dir / self.TARGET_FORMAT / rel_path.parent / f"{source_file.stem}.docx"
        return output_path
    
    def convert(self, source_file: Path) -> ConversionResult:
        """Convert TXT to DOCX with formatting."""
        start_time = datetime.now()
        result = ConversionResult(
            source_file=source_file,
            target_format=self.TARGET_FORMAT
        )
        
        try:
            output_file = self.get_output_path(source_file)
            
            if self._should_skip(source_file, output_file):
                result.error = f"Skipped: output exists and is up-to-date"
                result.output_file = output_file
                return result
            
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'Text Document', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f'Source: {source_file.name}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('')
            
            # Read TXT content
            with open(source_file, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Split by lines and add as paragraphs
            lines = content.split('\n')
            for line in lines:
                if line.strip():  # Skip empty lines
                    p = doc.add_paragraph(line)
                    # Set font
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Courier New'  # Monospace for technical text
                else:
                    doc.add_paragraph('')  # Preserve empty lines
            
            # Save document
            doc.save(output_file)
            
            result.success = True
            result.output_file = output_file
            result.output_size_bytes = output_file.stat().st_size
            
            self.logger.info(f"✓ TXT → DOCX: {source_file.name} ({result.output_size_bytes / 1024:.1f} KB)")
            
        except Exception as e:
            result.error = str(e)
            self.logger.error(f"✗ TXT → DOCX failed: {source_file.name}: {e}")
        
        finally:
            result.duration_sec = (datetime.now() - start_time).total_seconds()
        
        return result


class CSVToJSONConverter(FileConverter):
    """Converts CSV files to JSON format."""
    
    TARGET_FORMAT = "json"
    
    def can_convert(self, source_file: Path) -> bool:
        return source_file.suffix.lower() == ".csv"
    
    def get_output_path(self, source_file: Path) -> Path:
        rel_path = source_file.relative_to(self.config.source_dir)
        output_path = self.config.output_dir / self.TARGET_FORMAT / rel_path.parent / f"{source_file.stem}.json"
        return output_path
    
    def convert(self, source_file: Path) -> ConversionResult:
        """Convert CSV to JSON array of objects."""
        start_time = datetime.now()
        result = ConversionResult(
            source_file=source_file,
            target_format=self.TARGET_FORMAT
        )
        
        try:
            output_file = self.get_output_path(source_file)
            
            if self._should_skip(source_file, output_file):
                result.error = f"Skipped: output exists and is up-to-date"
                result.output_file = output_file
                return result
            
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Read CSV
            with open(source_file, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                data = list(reader)
            
            # Write JSON
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            
            result.success = True
            result.output_file = output_file
            result.output_size_bytes = output_file.stat().st_size
            
            self.logger.info(f"✓ CSV → JSON: {source_file.name} ({result.output_size_bytes / 1024:.1f} KB)")
            
        except Exception as e:
            result.error = str(e)
            self.logger.error(f"✗ CSV → JSON failed: {source_file.name}: {e}")
        
        finally:
            result.duration_sec = (datetime.now() - start_time).total_seconds()
        
        return result


class XMLToJSONConverter(FileConverter):
    """Converts XML files to JSON format."""
    
    TARGET_FORMAT = "json"
    
    def can_convert(self, source_file: Path) -> bool:
        return source_file.suffix.lower() == ".xml"
    
    def get_output_path(self, source_file: Path) -> Path:
        rel_path = source_file.relative_to(self.config.source_dir)
        output_path = self.config.output_dir / self.TARGET_FORMAT / rel_path.parent / f"{source_file.stem}_xml.json"
        return output_path
    
    def convert(self, source_file: Path) -> ConversionResult:
        """Convert XML to JSON preserving structure."""
        start_time = datetime.now()
        result = ConversionResult(
            source_file=source_file,
            target_format=self.TARGET_FORMAT
        )
        
        try:
            output_file = self.get_output_path(source_file)
            
            if self._should_skip(source_file, output_file):
                result.error = f"Skipped: output exists and is up-to-date"
                result.output_file = output_file
                return result
            
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Parse XML
            tree = ET.parse(source_file)
            root = tree.getroot()
            
            # Convert to dict
            data = self._xml_to_dict(root)
            
            # Write JSON
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            
            result.success = True
            result.output_file = output_file
            result.output_size_bytes = output_file.stat().st_size
            
            self.logger.info(f"✓ XML → JSON: {source_file.name} ({result.output_size_bytes / 1024:.1f} KB)")
            
        except Exception as e:
            result.error = str(e)
            self.logger.error(f"✗ XML → JSON failed: {source_file.name}: {e}")
        
        finally:
            result.duration_sec = (datetime.now() - start_time).total_seconds()
        
        return result
    
    def _xml_to_dict(self, element: ET.Element) -> Dict:
        """Recursively convert XML element to dictionary."""
        result = {}
        
        # Add attributes
        if element.attrib:
            result["@attributes"] = dict(element.attrib)
        
        # Add text content
        if element.text and element.text.strip():
            result["@text"] = element.text.strip()
        
        # Add children
        for child in element:
            child_data = self._xml_to_dict(child)
            
            if child.tag in result:
                # Multiple children with same tag -> convert to list
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_data)
            else:
                result[child.tag] = child_data
        
        # If only text and no attributes/children, return text directly
        if len(result) == 1 and "@text" in result:
            return result["@text"]
        
        return result


class PDFToImagesPDFConverter(FileConverter):
    """Converts PDF with text to images and back to image-based PDF (for OCR testing)."""
    
    TARGET_FORMAT = "pdf_images"
    
    def can_convert(self, source_file: Path) -> bool:
        return source_file.suffix.lower() == ".pdf" and PDF2IMAGE_AVAILABLE
    
    def get_output_path(self, source_file: Path) -> Path:
        """Return PDF path for image-based PDF."""
        rel_path = source_file.relative_to(self.config.source_dir)
        output_file = self.config.output_dir / self.TARGET_FORMAT / rel_path.parent / f"{source_file.stem}_images.pdf"
        return output_file
    
    def convert(self, source_file: Path) -> ConversionResult:
        """Convert PDF to images and back to image-based PDF."""
        start_time = datetime.now()
        result = ConversionResult(
            source_file=source_file,
            target_format=self.TARGET_FORMAT
        )
        
        try:
            output_file = self.get_output_path(source_file)
            
            # Check if should skip
            if self._should_skip(source_file, output_file):
                result.error = f"Skipped: output exists and is up-to-date"
                result.output_file = output_file
                return result
            
            # Create output directory
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Convert PDF to images
            images = convert_from_path(
                source_file,
                dpi=self.config.image_dpi,
                fmt=self.config.image_format
            )
            
            # Save images back to PDF (image-based, no text layer)
            if images:
                # Convert PIL images to RGB mode for PDF
                rgb_images = [img.convert('RGB') for img in images]
                
                # Save as PDF
                rgb_images[0].save(
                    output_file,
                    save_all=True,
                    append_images=rgb_images[1:] if len(rgb_images) > 1 else [],
                    resolution=self.config.image_dpi
                )
            
            result.success = True
            result.output_file = output_file
            result.output_size_bytes = output_file.stat().st_size
            
            self.logger.info(f"✓ PDF → Images → PDF: {source_file.name} ({len(images)} pages, {result.output_size_bytes / 1024:.1f} KB)")
            
        except Exception as e:
            result.error = str(e)
            self.logger.error(f"✗ PDF → Images → PDF failed: {source_file.name}: {e}")
        
        finally:
            result.duration_sec = (datetime.now() - start_time).total_seconds()
        
        return result


# =============================================================================
# CONVERSION ORCHESTRATOR
# =============================================================================

class DatasetConverter:
    """Main orchestrator for dataset conversion."""
    
    def __init__(self, config: ConversionConfig):
        self.config = config
        self.logger = self._setup_logger()
        self.converters = self._initialize_converters()
        self.summary = ConversionSummary()
    
    def _setup_logger(self) -> logging.Logger:
        """Setup logging configuration."""
        logger = logging.getLogger("DatasetConverter")
        logger.setLevel(logging.DEBUG if self.config.verbose else logging.INFO)
        
        # Console handler
        handler = logging.StreamHandler()
        handler.setLevel(logging.DEBUG if self.config.verbose else logging.INFO)
        
        formatter = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%H:%M:%S'
        )
        handler.setFormatter(formatter)
        
        logger.addHandler(handler)
        
        # File handler
        log_file = self.config.output_dir / "conversion.log"
        log_file.parent.mkdir(parents=True, exist_ok=True)
        
        file_handler = logging.FileHandler(log_file, mode='w')
        file_handler.setLevel(logging.DEBUG)
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        
        return logger
    
    def _initialize_converters(self) -> List[FileConverter]:
        """Initialize all available converters."""
        converters = []
        
        # Check dependencies and create converters - SPECIFIC CONVERSIONS ONLY
        if "xlsx" in self.config.formats:
            if OPENPYXL_AVAILABLE:
                converters.append(CSVToXLSXConverter(self.config, self.logger))
            else:
                self.logger.warning("XLSX conversion disabled: openpyxl not installed")
        
        if "docx" in self.config.formats:
            if DOCX_AVAILABLE:
                converters.append(TXTToDOCXConverter(self.config, self.logger))
            else:
                self.logger.warning("DOCX conversion disabled: python-docx not installed")
        
        if "json" in self.config.formats:
            converters.append(XMLToJSONConverter(self.config, self.logger))
        
        if "pdf_images" in self.config.formats:
            if PDF2IMAGE_AVAILABLE:
                converters.append(PDFToImagesPDFConverter(self.config, self.logger))
            else:
                self.logger.warning("PDF → Images → PDF conversion disabled: pdf2image not installed")
        
        return converters
    
    def collect_source_files(self) -> List[Path]:
        """Collect all source files to convert."""
        files = []
        
        # Scan for CSV, TXT, XML, and PDF files (specific conversions)
        for pattern in ["**/*.csv", "**/*.txt", "**/*.xml", "**/*.pdf"]:
            files.extend(self.config.source_dir.glob(pattern))
        
        # Filter out already converted files (e.g., .anonymous files)
        files = [f for f in files if not any(
            suffix in f.suffixes for suffix in [".anonymous", ".anon"]
        )]
        
        self.logger.info(f"Found {len(files)} source files to process")
        return sorted(files)
    
    def convert_file(self, source_file: Path) -> List[ConversionResult]:
        """Convert a single file to all applicable formats."""
        results = []
        
        for converter in self.converters:
            if converter.can_convert(source_file):
                result = converter.convert(source_file)
                results.append(result)
                self.summary.add_result(result)
        
        return results
    
    def run(self):
        """Run the conversion process."""
        self.logger.info("="*70)
        self.logger.info("DATASET CONVERSION STARTED")
        self.logger.info("="*70)
        self.logger.info(f"Source: {self.config.source_dir}")
        self.logger.info(f"Output: {self.config.output_dir}")
        self.logger.info(f"Formats: {', '.join(sorted(self.config.formats))}")
        self.logger.info(f"Workers: {self.config.max_workers}")
        self.logger.info("")
        
        # Check dependencies
        self._check_dependencies()
        
        # Collect files
        source_files = self.collect_source_files()
        if not source_files:
            self.logger.error("No source files found!")
            return
        
        self.summary.total_files = len(source_files) * len(self.converters)
        
        # Convert files
        if self.config.max_workers > 1:
            self._convert_parallel(source_files)
        else:
            self._convert_sequential(source_files)
        
        # Finalize and print summary
        self.summary.finalize()
        self._print_summary()
    
    def _convert_sequential(self, source_files: List[Path]):
        """Convert files sequentially."""
        for i, source_file in enumerate(source_files, 1):
            self.logger.info(f"[{i}/{len(source_files)}] Processing: {source_file.name}")
            self.convert_file(source_file)
    
    def _convert_parallel(self, source_files: List[Path]):
        """Convert files in parallel."""
        with ThreadPoolExecutor(max_workers=self.config.max_workers) as executor:
            futures = {
                executor.submit(self.convert_file, f): f 
                for f in source_files
            }
            
            for i, future in enumerate(as_completed(futures), 1):
                source_file = futures[future]
                try:
                    future.result()
                    self.logger.info(f"[{i}/{len(source_files)}] Completed: {source_file.name}")
                except Exception as e:
                    self.logger.error(f"[{i}/{len(source_files)}] Error processing {source_file.name}: {e}")
    
    def _check_dependencies(self):
        """Check and log dependency availability."""
        self.logger.info("Dependency Check:")
        deps = {
            "openpyxl": OPENPYXL_AVAILABLE,
            "python-docx": DOCX_AVAILABLE,
            "pdf2image": PDF2IMAGE_AVAILABLE,
            "pandas": PANDAS_AVAILABLE,
        }
        
        for dep, available in deps.items():
            status = "✓" if available else "✗"
            self.logger.info(f"  {status} {dep}")
        
        self.logger.info("")
    
    def _print_summary(self):
        """Print conversion summary."""
        self.logger.info("")
        self.logger.info("="*70)
        self.logger.info("CONVERSION SUMMARY")
        self.logger.info("="*70)
        self.logger.info(f"Total files processed: {len(self.summary.results)}")
        self.logger.info(f"Successful: {self.summary.successful}")
        self.logger.info(f"Failed: {self.summary.failed}")
        self.logger.info(f"Skipped: {self.summary.skipped}")
        self.logger.info(f"Duration: {self.summary.duration_sec:.1f} seconds")
        
        # Group by format
        by_format: Dict[str, List[ConversionResult]] = {}
        for result in self.summary.results:
            if result.success:
                by_format.setdefault(result.target_format, []).append(result)
        
        self.logger.info("")
        self.logger.info("By Format:")
        for fmt, results in sorted(by_format.items()):
            total_size = sum(r.output_size_bytes for r in results)
            self.logger.info(f"  {fmt.upper()}: {len(results)} files ({total_size / (1024*1024):.1f} MB)")
        
        # Show failures
        failures = [r for r in self.summary.results if not r.success and "skipped" not in (r.error or "").lower()]
        if failures:
            self.logger.info("")
            self.logger.info("Failures:")
            for result in failures[:10]:  # Show first 10
                self.logger.error(f"  {result.source_file.name} → {result.target_format}: {result.error}")
            if len(failures) > 10:
                self.logger.error(f"  ... and {len(failures) - 10} more")
        
        self.logger.info("")
        self.logger.info(f"Log file: {self.config.output_dir / 'conversion.log'}")
        self.logger.info("="*70)


# =============================================================================
# CLI
# =============================================================================

def create_parser() -> argparse.ArgumentParser:
    """Create argument parser."""
    parser = argparse.ArgumentParser(
        description="Convert vulnnet_scans_openvas dataset to multiple formats",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert to all formats
  python convert_dataset.py --all
  
  # Convert to specific formats
  python convert_dataset.py --formats xlsx docx json
  
  # Convert with custom directories
  python convert_dataset.py --source /path/to/data --output /path/to/output --all
  
  # Parallel conversion with 8 workers
  python convert_dataset.py --all --workers 8
  
  # Force reconversion (ignore existing files)
  python convert_dataset.py --all --no-skip-existing
        """
    )
    
    parser.add_argument(
        "--source", type=str,
        default="vulnnet_scans_openvas",
        help="Source directory (default: vulnnet_scans_openvas)"
    )
    
    parser.add_argument(
        "--output", type=str,
        default="benchmark/converted_datasets",
        help="Output directory (default: benchmark/converted_datasets)"
    )
    
    parser.add_argument(
        "--formats", nargs="+",
        choices=["xlsx", "docx", "json", "pdf_images"],
        help="Formats to generate: xlsx (CSV→XLSX), docx (TXT→DOCX), json (XML→JSON), pdf_images (PDF→Images→PDF)"
    )
    
    parser.add_argument(
        "--all", action="store_true",
        help="Convert to all formats (CSV→XLSX, TXT→DOCX, XML→JSON, PDF→Images→PDF)"
    )
    
    parser.add_argument(
        "--workers", type=int, default=4,
        help="Number of parallel workers (default: 4)"
    )
    
    parser.add_argument(
        "--no-skip-existing", action="store_true",
        help="Reconvert even if output exists"
    )
    
    parser.add_argument(
        "--image-dpi", type=int, default=150,
        help="DPI for PDF to image conversion (default: 150)"
    )
    
    parser.add_argument(
        "--image-format", type=str, default="png",
        choices=["png", "jpg", "jpeg"],
        help="Image format for PDF conversion (default: png)"
    )
    
    parser.add_argument(
        "--verbose", "-v", action="store_true",
        help="Verbose output"
    )
    
    return parser


def main():
    """Main entry point."""
    parser = create_parser()
    args = parser.parse_args()
    
    # Determine formats
    if args.all:
        formats = {"xlsx", "docx", "json", "pdf_images"}
    elif args.formats:
        formats = set(args.formats)
    else:
        print("Error: specify --all or --formats")
        parser.print_help()
        return 1
    
    # Create configuration
    try:
        config = ConversionConfig(
            source_dir=Path(args.source),
            output_dir=Path(args.output),
            formats=formats,
            max_workers=args.workers,
            skip_existing=not args.no_skip_existing,
            image_dpi=args.image_dpi,
            image_format=args.image_format,
            verbose=args.verbose
        )
    except ValueError as e:
        print(f"Configuration error: {e}")
        return 1
    
    # Run conversion
    converter = DatasetConverter(config)
    converter.run()
    
    return 0 if converter.summary.failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
