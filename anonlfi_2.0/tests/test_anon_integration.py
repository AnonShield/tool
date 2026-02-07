import unittest
import os
import subprocess
import shutil
import json
import openpyxl
import fitz
from docx import Document
import sqlite3
import datetime
from PIL import Image, ImageDraw, ImageFont
import re

class TestAnonIntegration(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.test_data_dir = "tests/test_data"
        cls.output_dir = "output"
        cls.db_path = os.path.join("db", "entities.db")

        if os.path.exists(cls.test_data_dir):
            shutil.rmtree(cls.test_data_dir)
        os.makedirs(cls.test_data_dir, exist_ok=True)
        
        if os.path.exists(cls.output_dir):
            shutil.rmtree(cls.output_dir)
        os.makedirs(cls.output_dir, exist_ok=True)

        os.makedirs(os.path.dirname(cls.db_path), exist_ok=True)
        if os.path.exists(cls.db_path):
            os.remove(cls.db_path)

        cls.text1 = "My name is John Doe."
        cls.image_text = "My email is secret.email@example.com."
        cls.text2 = "I live in New York."
        
        email = "test@example.com"
        text = f"My name is John Doe and my email is {email}."

        with open(os.path.join(cls.test_data_dir, "test.txt"    ), "w") as f:
            f.write(text)

        with open(os.path.join(cls.test_data_dir, "test.csv"), "w") as f:
            f.write(f"name,email\nJohn Doe,{email}\nJane Smith,jane.smith@example.com")

        with open(os.path.join(cls.test_data_dir, "test.json"), "w") as f:
            json.dump({"user": {"name": "John Doe", "email": email}}, f)

        with open(os.path.join(cls.test_data_dir, "test.xml"), "w") as f:
            f.write(f"<user>\n  <name>John Doe</name>\n  <email>{email}</email>\n</user>")

        img = Image.new('RGB', (1200, 200), color = (255, 255, 255))
        d = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 40)
        except IOError:
            font = ImageFont.load_default()
        d.text((10,10), cls.image_text, fill=(0,0,0), font=font)
        
        cls.image_path = os.path.join(cls.test_data_dir, "test_image.png")
        img.save(cls.image_path, "PNG")
        
        pdf_path = os.path.join(cls.test_data_dir, "test.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 72), cls.text1)
        page.insert_image(fitz.Rect(50, 100, 450, 200), filename=cls.image_path)
        page.insert_text((50, 220), cls.text2)
        doc.save(pdf_path)
        doc.close()

        docx_path = os.path.join(cls.test_data_dir, "test.docx")
        doc = Document()
        doc.add_paragraph(cls.text1)
        doc.add_picture(cls.image_path)
        doc.add_paragraph(cls.text2)
        doc.save(docx_path)

        xlsx_path = os.path.join(cls.test_data_dir, "test.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.cell(row=1, column=1, value="name")
        ws.cell(row=1, column=2, value="email")
        ws.cell(row=2, column=1, value="John Doe")
        ws.cell(row=2, column=2, value=email)
        wb.save(xlsx_path)

    @classmethod
    def tearDownClass(cls):
        if os.path.exists(cls.test_data_dir):
            shutil.rmtree(cls.test_data_dir)
        if os.path.exists(cls.output_dir):
            shutil.rmtree(cls.output_dir)
        if os.path.exists(cls.db_path):
            os.remove(cls.db_path)
        if os.path.exists(os.path.dirname(cls.db_path)) and not os.listdir(os.path.dirname(cls.db_path)):
            os.rmdir(os.path.dirname(cls.db_path))

    def setUp(self):
        self.secret_key = "test-secret-key"
        os.environ["ANON_SECRET_KEY"] = self.secret_key
        if os.path.exists(self.output_dir):
            shutil.rmtree(self.output_dir)
        os.makedirs(self.output_dir, exist_ok=True)

    def _run_anon_py(self, path, lang="en", preserve_entities="", allow_list="", slug_length=None):
        cmd = [
            "python",
            "anon.py",
            path,
            "--lang",
            lang,
        ]
        if preserve_entities:
            cmd.extend(["--preserve-entities", preserve_entities])
        if allow_list:
            cmd.extend(["--allow-list", allow_list])
        if slug_length is not None:
            cmd.extend(["--slug-length", str(slug_length)])
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            print("anon.py failed for {}:".format(path))
            print("STDOUT: {}".format(result.stdout))
            print("STDERR: {}".format(result.stderr))
        self.assertEqual(result.returncode, 0)
        return result

    def _get_output_file_path(self, input_file_path):
        base_name, input_ext = os.path.splitext(os.path.basename(input_file_path))
        if input_ext.lower() in [".csv", ".xlsx", ".json", ".xml"]:
            output_ext = input_ext.lower()
        else:
            output_ext = ".txt"
        return os.path.join(self.output_dir, "anon_{}{}".format(base_name, output_ext))

    def test_txt_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.txt")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertNotIn("John Doe", content)
        self.assertNotIn("test@example.com", content)

    def test_csv_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.csv")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertNotIn("John Doe", content)
        self.assertNotIn("test@example.com", content)
        self.assertNotIn("Jane Smith", content)
        self.assertNotIn("jane.smith@example.com", content)

    def test_json_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.json")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            data = json.load(f)
        
        self.assertIn("[PERSON_", data["user"]["name"])
        self.assertIn("[EMAIL_ADDRESS_", data["user"]["email"])
        self.assertNotIn("John Doe", data["user"]["name"])
        self.assertNotIn("test@example.com", data["user"]["email"])

    def test_xml_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.xml")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()

        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertNotIn("John Doe", content)
        self.assertNotIn("test@example.com", content)

    def test_docx_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.docx")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertIn("[LOCATION_", content)
        self.assertNotIn("John Doe", content)
        self.assertNotIn("secret.email@example.com", content)
        self.assertNotIn("New York", content)

        person_index = content.find("[PERSON_")
        email_index = content.find("[EMAIL_ADDRESS_")
        location_index = content.find("[LOCATION_")

        self.assertTrue(person_index != -1 and email_index != -1 and location_index != -1)
        self.assertTrue(person_index < email_index < location_index)

    def test_xlsx_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.xlsx")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        wb = openpyxl.load_workbook(output_file)
        ws = wb.active
        self.assertIn("[PERSON_", ws.cell(row=2, column=1).value)
        self.assertIn("[EMAIL_ADDRESS_", ws.cell(row=2, column=2).value)
        self.assertNotIn("John Doe", ws.cell(row=2, column=1).value)
        self.assertNotIn("test@example.com", ws.cell(row=2, column=2).value)

    def test_pdf_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.pdf")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertIn("[LOCATION_", content)
        self.assertNotIn("John Doe", content)
        self.assertNotIn("secret.email@example.com", content)
        self.assertNotIn("New York", content)

        person_index = content.find("[PERSON_")
        email_index = content.find("[EMAIL_ADDRESS_")
        location_index = content.find("[LOCATION_")

        self.assertTrue(person_index != -1 and email_index != -1 and location_index != -1)
        self.assertTrue(person_index < email_index < location_index)

    def test_image_format_anonymization(self):
        image_formats = {
            "png": "PNG",
            "jpeg": "JPEG",
            "gif": "GIF",
            "bmp": "BMP",
            "tiff": "TIFF",
            "webp": "WEBP",
            "jp2": "JPEG2000",
        }
        image_text = "This image contains an email: image.test@example.com."

        img = Image.new('RGB', (1200, 200), color = (255, 255, 255))
        d = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 40)
        except IOError:
            font = ImageFont.load_default()
        d.text((10,10), image_text, fill=(0,0,0), font=font)

        for ext, pillow_format in image_formats.items():
            with self.subTest(ext=ext):
                file_name = f"test_image_format.{ext}"
                image_path = os.path.join(self.test_data_dir, file_name)
                
                try:
                    img.save(image_path, pillow_format)
                except (KeyError, ValueError) as e:
                    self.skipTest(f"Saving with Pillow format {pillow_format} for .{ext} failed: {e}")
                    continue

                self._run_anon_py(image_path)
                output_file = self._get_output_file_path(image_path)

                self.assertTrue(os.path.exists(output_file), f"Output file for {ext} not created.")
                with open(output_file, "r") as f:
                    content = f.read()

                self.assertIn("[EMAIL_ADDRESS_", content)
                self.assertNotIn("image.test@example.com", content)

    def test_preserve_entities(self):
        test_file = os.path.join(self.test_data_dir, "test.txt")
        self._run_anon_py(test_file, preserve_entities="PERSON")
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertNotIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertIn("John Doe", content)
        self.assertNotIn("test@example.com", content)

    def test_allow_list(self):
        test_file = os.path.join(self.test_data_dir, "test.txt")
        self._run_anon_py(test_file, allow_list="John Doe")
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertNotIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertIn("John Doe", content)
        self.assertNotIn("test@example.com", content)

    def test_directory_anonymization(self):
        temp_dir = os.path.join(self.test_data_dir, "temp_dir_for_parallel_test")
        os.makedirs(temp_dir, exist_ok=True)

        shutil.copy(os.path.join(self.test_data_dir, "test.txt"), temp_dir)
        shutil.copy(os.path.join(self.test_data_dir, "test.csv"), temp_dir)
        shutil.copy(os.path.join(self.test_data_dir, "test.json"), temp_dir)

        self._run_anon_py(temp_dir)

        expected_output_txt = self._get_output_file_path(os.path.join(temp_dir, "test.txt"))
        expected_output_csv = self._get_output_file_path(os.path.join(temp_dir, "test.csv"))
        expected_output_json = self._get_output_file_path(os.path.join(temp_dir, "test.json"))

        self.assertTrue(os.path.exists(expected_output_txt))
        self.assertTrue(os.path.exists(expected_output_csv))
        self.assertTrue(os.path.exists(expected_output_json))

        with open(expected_output_txt, "r") as f:
            content = f.read()
        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)

        shutil.rmtree(temp_dir)

    def test_database_batching(self):
        if os.path.exists(self.db_path):
            os.remove(self.db_path)

        test_file = os.path.join(self.test_data_dir, "test.txt")
        self._run_anon_py(test_file)

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT first_seen, last_seen FROM entities")
        results = cursor.fetchall()
        conn.close()

        self.assertGreater(len(results), 0, "No entities found in the database.")

        for first_seen_str, last_seen_str in results:
            first_seen = datetime.datetime.fromisoformat(first_seen_str)
            last_seen = datetime.datetime.fromisoformat(last_seen_str)
            self.assertLessEqual((last_seen - first_seen).total_seconds(), 1)

    def test_list_languages(self):
        cmd = ["python", "anon.py", "--list-languages"]
        result = subprocess.run(cmd, capture_output=True, text=True)
        self.assertEqual(result.returncode, 0)
        self.assertIn("[*] Supported languages:", result.stdout)
        self.assertIn("en: English", result.stdout)
        self.assertIn("pt: Portuguese", result.stdout)

    def test_slug_length(self):
        if os.path.exists(self.db_path):
            os.remove(self.db_path)

        test_file = os.path.join(self.test_data_dir, "test.txt")
        slug_len = 8
        self._run_anon_py(test_file, slug_length=slug_len)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        match = re.search(r"\[[A-Z_]+_([a-f0-9]+)\]", content)
        self.assertIsNotNone(match, "No slug found in anonymized content.")
        extracted_slug = match.group(1)
        self.assertEqual(len(extracted_slug), slug_len, f"Expected slug length {slug_len}, got {len(extracted_slug)}.")

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT slug_name FROM entities WHERE slug_name = ?", (extracted_slug,))
        db_slug = cursor.fetchone()
        conn.close()
        self.assertIsNotNone(db_slug, "Slug not found in database.")
        self.assertEqual(len(db_slug[0]), slug_len, f"Expected database slug length {slug_len}, got {len(db_slug[0])}.")

if __name__ == "__main__":
    unittest.main()
