import unittest
import os
import subprocess
import shutil
import json
import openpyxl
import fitz
from docx import Document
import sqlite3
import datetime
from PIL import Image, ImageDraw, ImageFont
import re
import time

class TestAnonIntegration(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.test_data_dir = "tests/test_data_integration"
        cls.output_dir = "test_output_integration"
        cls.db_dir = "db_pytest_integration"
        cls.db_path = os.path.join(cls.db_dir, "entities.db")

        # Clean up and create directories
        for d in [cls.test_data_dir, cls.output_dir, cls.db_dir]:
            if os.path.exists(d):
                shutil.rmtree(d)
            os.makedirs(d, exist_ok=True)

        cls.text1 = "My name is John Doe."
        cls.image_text = "My email is secret.email@example.com."
        cls.text2 = "I live in New York."
        
        email = "test@example.com"
        text = f"My name is John Doe and my email is {email}."

        with open(os.path.join(cls.test_data_dir, "test.txt"), "w") as f:
            f.write(text)

        with open(os.path.join(cls.test_data_dir, "test.csv"), "w") as f:
            f.write(f"name,email\nJohn Doe,{email}\nJane Smith,jane.smith@example.com")

        with open(os.path.join(cls.test_data_dir, "test.json"), "w") as f:
            json.dump({"user": {"name": "John Doe", "email": email}}, f)

        with open(os.path.join(cls.test_data_dir, "test.xml"), "w") as f:
            f.write(f'<?xml version="1.0" encoding="UTF-8"?>\n<user>\n  <name>John Doe</name>\n  <email>{email}</email>\n</user>')

        img = Image.new('RGB', (1200, 200), color = (255, 255, 255))
        d = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 40)
        except IOError:
            font = ImageFont.load_default()
        d.text((10,10), cls.image_text, fill=(0,0,0), font=font)
        
        cls.image_path = os.path.join(cls.test_data_dir, "test_image.png")
        img.save(cls.image_path, "PNG")
        
        pdf_path = os.path.join(cls.test_data_dir, "test.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 72), cls.text1)
        page.insert_image(fitz.Rect(50, 100, 450, 200), filename=cls.image_path)
        page.insert_text((50, 220), cls.text2)
        doc.save(pdf_path)
        doc.close()

        docx_path = os.path.join(cls.test_data_dir, "test.docx")
        doc = Document()
        doc.add_paragraph(cls.text1)
        doc.add_picture(cls.image_path)
        doc.add_paragraph(cls.text2)
        doc.save(docx_path)

        xlsx_path = os.path.join(cls.test_data_dir, "test.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.cell(row=1, column=1, value="name")
        ws.cell(row=1, column=2, value="email")
        ws.cell(row=2, column=1, value="John Doe")
        ws.cell(row=2, column=2, value=email)
        wb.save(xlsx_path)

    @classmethod
    def tearDownClass(cls):
        for d in [cls.test_data_dir, cls.output_dir, cls.db_dir]:
            if os.path.exists(d):
                shutil.rmtree(d)

    def setUp(self):
        self.secret_key = "test-secret-key"
        os.environ["ANON_SECRET_KEY"] = self.secret_key
        # Clear DB file before each test
        if os.path.exists(self.db_path):
            os.remove(self.db_path)

    def _run_anon_py(self, path, lang="en", preserve_entities="", allow_list="", slug_length=None, anonymization_strategy="presidio", extra_args=None):
        cmd = [
            "python",
            os.path.join(os.getcwd(), "anon.py"),
            path,
            "--lang", lang,
            "--db-dir", self.db_dir,
            "--output-dir", self.output_dir,
        ]
        if preserve_entities:
            cmd.extend(["--preserve-entities", preserve_entities])
        if allow_list:
            cmd.extend(["--allow-list", allow_list])
        if slug_length is not None:
            cmd.extend(["--slug-length", str(slug_length)])
        if anonymization_strategy:
            cmd.extend(["--anonymization-strategy", anonymization_strategy])
        if extra_args:
            cmd.extend(extra_args)
        
        # Set up environment variables for the subprocess
        env = os.environ.copy()
        current_dir = os.getcwd()
        if "PYTHONPATH" in env:
            env["PYTHONPATH"] = current_dir + os.pathsep + env["PYTHONPATH"]
        else:
            env["PYTHONPATH"] = current_dir

        result = subprocess.run(cmd, capture_output=True, text=True, env=env, cwd=os.getcwd())
        print("anon.py stdout:\n", result.stdout)
        print("anon.py stderr:\n", result.stderr)
        # self.assertEqual(result.returncode, 0, 
        #                  f"anon.py failed for {path}!\n"
        #                  f"STDOUT:\n{result.stdout}\n"
        #                  f"STDERR:\n{result.stderr}")
        return result

    def _get_output_file_path(self, input_file_path):
        base_name, input_ext = os.path.splitext(os.path.basename(input_file_path))
        if input_ext.lower() in [".csv", ".xlsx", ".json", ".xml", ".jsonl"]:
            output_ext = input_ext.lower()
        else:
            output_ext = ".txt"
        return os.path.join(self.output_dir, "anon_{}{}".format(base_name, output_ext))

    def test_txt_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.txt")
        self._run_anon_py(test_file, extra_args=["--overwrite"])
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertNotIn("John Doe", content)
        self.assertNotIn("test@example.com", content)

    def test_csv_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.csv")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)
        self.assertNotIn("John Doe", content)
        self.assertNotIn("test@example.com", content)
        self.assertNotIn("Jane Smith", content)
        self.assertNotIn("jane.smith@example.com", content)

    def test_json_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.json")
        self._run_anon_py(test_file)
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            data = json.load(f)
        
        self.assertIn("[PERSON_", data["user"]["name"])
        self.assertIn("[EMAIL_ADDRESS_", data["user"]["email"])

    def test_xml_anonymization(self):
        test_file = os.path.join(self.test_data_dir, "test.xml")
        self._run_anon_py(test_file, extra_args=["--force-large-xml", "--overwrite"])
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()

        self.assertIn("[PERSON_", content)
        self.assertIn("[EMAIL_ADDRESS_", content)

    def test_database_batching(self):
        # This test relies on the subprocess writing to the DB
        test_file = os.path.join(self.test_data_dir, "test.txt")
        self._run_anon_py(test_file, extra_args=["--overwrite"])

        # The shutdown process should be blocking, so no sleep is needed.
        self.assertTrue(os.path.exists(self.db_path))
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM entities")
        results = cursor.fetchall()
        conn.close()

        self.assertGreater(len(results), 0, "No entities found in the database.")

    def test_slug_length(self):
        test_file = os.path.join(self.test_data_dir, "test.txt")
        slug_len = 8
        self._run_anon_py(test_file, slug_length=slug_len, extra_args=["--overwrite"])
        output_file = self._get_output_file_path(test_file)
        
        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()
        
        match = re.search(r"[[A-Z_]+_([a-f0-9]+)]", content)
        self.assertIsNotNone(match, "No slug found in anonymized content.")
        extracted_slug = match.group(1)
        self.assertEqual(len(extracted_slug), slug_len)

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT slug_name FROM entities WHERE slug_name = ?", (extracted_slug,))
        db_slug = cursor.fetchone()
        conn.close()
        self.assertIsNotNone(db_slug, "Slug not found in database.")
        self.assertEqual(len(db_slug[0]), slug_len)

    def test_slug_length_zero(self):
        """
        Tests the behavior of --slug-length 0.
        Ensures entities are anonymized without a slug in the text and are NOT saved to the database.
        This test uses a specific file and the 'fast' strategy to ensure the logic is tested
        on an entity type that this strategy can detect.
        """
        # This specific file only contains an email, which the 'fast' strategy can detect.
        test_file = os.path.join(self.test_data_dir, "test_slug_zero.txt")
        if not os.path.exists(test_file):
            with open(test_file, "w") as f:
                f.write("My email is test@example.com.")

        # Use fast strategy as it's the one we modified for this behavior.
        self._run_anon_py(test_file, slug_length=0, anonymization_strategy="fast", extra_args=["--overwrite"])
        output_file = self._get_output_file_path(test_file)

        self.assertTrue(os.path.exists(output_file))
        with open(output_file, "r") as f:
            content = f.read()

        # 1. Check that the placeholder is [ENTITY_TYPE] without a hash.
        self.assertIn("[EMAIL_ADDRESS]", content)
        self.assertNotIn("test@example.com", content)

        # 2. Check that no hash is present in the placeholder.
        self.assertNotIn("[EMAIL_ADDRESS_", content)
        
        # 3. Verify that NO entities were saved to the database.
        self.assertTrue(os.path.exists(self.db_path), "Database file should exist.")
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM entities")
        count = cursor.fetchone()[0]
        conn.close()

        self.assertEqual(count, 0, "Database should be empty when slug_length is 0.")

    def test_preserve_and_allow_list(self):
        test_file = os.path.join(self.test_data_dir, "test.txt")
        # The input is "My name is John Doe and my email is test@example.com."
        
        # 1. Test --preserve-entities
        self._run_anon_py(test_file, preserve_entities="PERSON", extra_args=["--overwrite"])
        output_file = self._get_output_file_path(test_file)
        with open(output_file, "r") as f:
            content = f.read()
        
        self.assertNotIn("[PERSON_", content, "PERSON should have been preserved.")
        self.assertIn("John Doe", content)
        self.assertIn("[EMAIL_ADDRESS_", content, "EMAIL_ADDRESS should have been anonymized.")
        self.assertNotIn("test@example.com", content)

        # 2. Test --allow-list
        self._run_anon_py(test_file, allow_list="test@example.com", extra_args=["--overwrite"])
        output_file = self._get_output_file_path(test_file)
        with open(output_file, "r") as f:
            content = f.read()

        self.assertIn("[PERSON_", content, "PERSON should have been anonymized.")
        self.assertNotIn("John Doe", content)
        self.assertNotIn("[EMAIL_ADDRESS_", content, "EMAIL_ADDRESS should have been on the allow list.")
        self.assertIn("test@example.com", content)

if __name__ == "__main__":
    unittest.main()
